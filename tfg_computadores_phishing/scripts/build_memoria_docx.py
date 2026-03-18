from __future__ import annotations

import argparse
from pathlib import Path
import re

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT_DIR = Path(__file__).resolve().parents[1]
DEFAULT_INPUT_MD = ROOT_DIR / "docs" / "MEMORIA_TFG_ETSII_APA7.md"
DEFAULT_OUTPUT_DOCX = ROOT_DIR / "docs" / "MEMORIA_TFG_ETSII_APA7.docx"
DEFAULT_REFERENCE_DOCX = ROOT_DIR / "docs" / "reference_urjc_tfg.docx"
DEFAULT_LOGO = ROOT_DIR / "docs" / "assets" / "branding" / "urjc_logo_oficial.png"


def _configure_margins(section) -> None:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.different_first_page_header_footer = False


def _set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for heading_name, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[heading_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def _clear_story(story) -> None:
    for child in list(story._element):
        story._element.remove(child)


def _apply_header_footer(section, logo_path: Path) -> None:
    header = section.header
    _clear_story(header)
    table = header.add_table(rows=1, cols=3, width=Cm(16))
    table.autofit = True
    left_cell, center_cell, right_cell = table.rows[0].cells

    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_left.add_run().add_picture(str(logo_path), width=Cm(2.1))

    p_center = center_cell.paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_center = p_center.add_run("Doble Grado en Ingenieria Informatica e Ingenieria de Computadores")
    r_center.font.name = "Times New Roman"
    r_center.font.size = Pt(9)
    r_center.font.bold = True

    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_right = p_right.add_run("CURSO 2024-2025")
    r_right.font.name = "Times New Roman"
    r_right.font.size = Pt(9)

    footer = section.footer
    _clear_story(footer)
    p_footer = footer.add_paragraph()
    _add_page_number(p_footer)


def _build_reference_doc(reference_docx: Path, logo_path: Path) -> None:
    doc = Document()
    _set_styles(doc)
    section = doc.sections[0]
    _configure_margins(section)
    _apply_header_footer(section, logo_path)
    reference_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(reference_docx))


def _postprocess_output(output_docx: Path, logo_path: Path) -> None:
    doc = Document(str(output_docx))
    _set_styles(doc)
    for section in doc.sections:
        _configure_margins(section)
        _apply_header_footer(section, logo_path)

    resumen_index = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().lower() == "resumen":
            resumen_index = idx
            break
    if resumen_index is not None:
        for idx, paragraph in enumerate(doc.paragraphs):
            if idx >= resumen_index:
                break
            if paragraph.text.strip():
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if resumen_index is not None and resumen_index > 0:
        previous = doc.paragraphs[resumen_index - 1]
        previous.add_run().add_break(WD_BREAK.PAGE)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name is None:
                run.font.name = "Times New Roman"
            if run.font.size is None and (paragraph.style and paragraph.style.name == "Normal"):
                run.font.size = Pt(12)

    doc.save(str(output_docx))


def _append_inline_markdown(paragraph, text: str) -> None:
    text = text.replace("**", "")
    text = text.replace("*", "")

    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def _flush_paragraph(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return

    paragraph = doc.add_paragraph()
    _append_inline_markdown(paragraph, " ".join(item.strip() for item in buffer))
    buffer.clear()


def _consume_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start

    while index < len(lines) and lines[index].lstrip().startswith("|"):
        raw = lines[index].strip().strip("|")
        cells = [cell.strip() for cell in raw.split("|")]
        rows.append(cells)
        index += 1

    return rows, index


def _is_separator_row(row: list[str]) -> bool:
    return all(cell.replace("-", "").replace(":", "").strip() == "" for cell in row)


def _write_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return

    header = rows[0]
    body = [row for row in rows[1:] if not _is_separator_row(row)]
    if not body:
        return

    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    for col_idx, value in enumerate(header):
        cell_paragraph = table.rows[0].cells[col_idx].paragraphs[0]
        run = cell_paragraph.add_run(value)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    for row in body:
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            cell_paragraph = cells[col_idx].paragraphs[0]
            _append_inline_markdown(cell_paragraph, value)


def _resolve_asset_path(input_md: Path, raw_path: str) -> Path:
    candidate = Path(raw_path.strip())
    if candidate.is_absolute():
        return candidate

    local_candidate = (input_md.parent / candidate).resolve()
    if local_candidate.exists():
        return local_candidate

    return (ROOT_DIR / candidate).resolve()


def _write_figure(doc: Document, input_md: Path, alt_text: str, raw_path: str) -> None:
    image_path = _resolve_asset_path(input_md, raw_path)
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if image_path.exists():
        image_paragraph.add_run().add_picture(str(image_path), width=Cm(15.5))
    else:
        missing_run = image_paragraph.add_run(f"[Figura no encontrada: {raw_path}]")
        missing_run.italic = True

    caption_text = alt_text.strip() if alt_text.strip() else image_path.name
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run(caption_text)
    caption_run.font.name = "Times New Roman"
    caption_run.font.size = Pt(11)
    caption_run.italic = True


def _build_docx_without_pandoc(input_md: Path, output_docx: Path, logo_path: Path) -> None:
    doc = Document()
    _set_styles(doc)
    section = doc.sections[0]
    _configure_margins(section)
    _apply_header_footer(section, logo_path)

    lines = input_md.read_text(encoding="utf-8-sig").splitlines()
    paragraph_buffer: list[str] = []
    index = 0

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()

        if not line:
            _flush_paragraph(doc, paragraph_buffer)
            index += 1
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading_match:
            _flush_paragraph(doc, paragraph_buffer)
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if level == 1:
                paragraph = doc.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _append_inline_markdown(paragraph, text)
            else:
                paragraph = doc.add_paragraph(style=f"Heading {level - 1}")
                _append_inline_markdown(paragraph, text)
            index += 1
            continue

        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", line)
        if image_match:
            _flush_paragraph(doc, paragraph_buffer)
            _write_figure(doc, input_md, image_match.group(1), image_match.group(2))
            index += 1
            continue

        if line.startswith("|"):
            _flush_paragraph(doc, paragraph_buffer)
            rows, index = _consume_table(lines, index)
            _write_table(doc, rows)
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", line)
        if ordered_match:
            _flush_paragraph(doc, paragraph_buffer)
            paragraph = doc.add_paragraph(style="List Number")
            _append_inline_markdown(paragraph, ordered_match.group(1).strip())
            index += 1
            continue

        unordered_match = re.match(r"^-\s+(.*)$", line)
        if unordered_match:
            _flush_paragraph(doc, paragraph_buffer)
            paragraph = doc.add_paragraph(style="List Bullet")
            _append_inline_markdown(paragraph, unordered_match.group(1).strip())
            index += 1
            continue

        quote_match = re.match(r"^>\s+(.*)$", line)
        if quote_match:
            _flush_paragraph(doc, paragraph_buffer)
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1)
            paragraph.paragraph_format.right_indent = Cm(1)
            _append_inline_markdown(paragraph, quote_match.group(1).strip())
            index += 1
            continue

        paragraph_buffer.append(line)
        index += 1

    _flush_paragraph(doc, paragraph_buffer)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))


def build_memoria(input_md: Path, output_docx: Path, reference_docx: Path, logo_path: Path) -> None:
    if not logo_path.exists():
        raise FileNotFoundError(f"Logo URJC no encontrado: {logo_path}")
    if not input_md.exists():
        raise FileNotFoundError(f"Markdown no encontrado: {input_md}")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    _build_reference_doc(reference_docx, logo_path)

    try:
        pypandoc.convert_file(
            str(input_md),
            "docx",
            outputfile=str(output_docx),
            extra_args=[
                f"--reference-doc={reference_docx}",
                f"--resource-path={ROOT_DIR / 'docs'}",
                "--standalone",
                "--toc",
                "--toc-depth=3",
            ],
        )
    except OSError:
        _build_docx_without_pandoc(input_md, output_docx, logo_path)

    _postprocess_output(output_docx, logo_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Genera la memoria DOCX con estilo URJC (portada + cabeceras).")
    parser.add_argument("--input-md", default=str(DEFAULT_INPUT_MD))
    parser.add_argument("--output-docx", default=str(DEFAULT_OUTPUT_DOCX))
    parser.add_argument("--reference-docx", default=str(DEFAULT_REFERENCE_DOCX))
    parser.add_argument("--logo", default=str(DEFAULT_LOGO))
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    build_memoria(
        input_md=Path(args.input_md),
        output_docx=Path(args.output_docx),
        reference_docx=Path(args.reference_docx),
        logo_path=Path(args.logo),
    )
    print(f"OK -> {args.output_docx}")


if __name__ == "__main__":
    main()
