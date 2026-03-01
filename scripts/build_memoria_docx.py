from __future__ import annotations

import argparse
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT_DIR = Path(__file__).resolve().parents[1]
DEFAULT_INPUT_MD = ROOT_DIR / "docs" / "MEMORIA_TFG_ETSII_APA7.md"
DEFAULT_OUTPUT_DOCX = ROOT_DIR / "docs" / "MEMORIA_TFG_ETSII_APA7.docx"
DEFAULT_REFERENCE_DOCX = ROOT_DIR / "docs" / "reference_urjc_tfg.docx"
DEFAULT_LOGO = ROOT_DIR / "docs" / "assets" / "branding" / "urjc_logo_oficial.png"


def _configure_margins(section) -> None:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.different_first_page_header_footer = False


def _set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for heading_name, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[heading_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def _clear_story(story) -> None:
    for child in list(story._element):
        story._element.remove(child)


def _apply_header_footer(section, logo_path: Path) -> None:
    header = section.header
    _clear_story(header)
    table = header.add_table(rows=1, cols=3, width=Cm(16))
    table.autofit = True
    left_cell, center_cell, right_cell = table.rows[0].cells

    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_left.add_run().add_picture(str(logo_path), width=Cm(2.1))

    p_center = center_cell.paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_center = p_center.add_run("Doble Grado en Ingenieria Informatica e Ingenieria de Computadores")
    r_center.font.name = "Times New Roman"
    r_center.font.size = Pt(9)
    r_center.font.bold = True

    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_right = p_right.add_run("CURSO 2024-2025")
    r_right.font.name = "Times New Roman"
    r_right.font.size = Pt(9)

    footer = section.footer
    _clear_story(footer)
    p_footer = footer.add_paragraph()
    _add_page_number(p_footer)


def _build_reference_doc(reference_docx: Path, logo_path: Path) -> None:
    doc = Document()
    _set_styles(doc)
    section = doc.sections[0]
    _configure_margins(section)
    _apply_header_footer(section, logo_path)
    reference_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(reference_docx))


def _postprocess_output(output_docx: Path, logo_path: Path) -> None:
    doc = Document(str(output_docx))
    _set_styles(doc)
    for section in doc.sections:
        _configure_margins(section)
        _apply_header_footer(section, logo_path)

    resumen_index = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().lower() == "resumen":
            resumen_index = idx
            break
    if resumen_index is not None:
        for idx, paragraph in enumerate(doc.paragraphs):
            if idx >= resumen_index:
                break
            if paragraph.text.strip():
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if resumen_index is not None and resumen_index > 0:
        previous = doc.paragraphs[resumen_index - 1]
        previous.add_run().add_break(WD_BREAK.PAGE)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name is None:
                run.font.name = "Times New Roman"
            if run.font.size is None and (paragraph.style and paragraph.style.name == "Normal"):
                run.font.size = Pt(12)

    doc.save(str(output_docx))


def build_memoria(input_md: Path, output_docx: Path, reference_docx: Path, logo_path: Path) -> None:
    if not logo_path.exists():
        raise FileNotFoundError(f"Logo URJC no encontrado: {logo_path}")
    if not input_md.exists():
        raise FileNotFoundError(f"Markdown no encontrado: {input_md}")

    _build_reference_doc(reference_docx, logo_path)
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    pypandoc.convert_file(
        str(input_md),
        "docx",
        outputfile=str(output_docx),
        extra_args=[
            f"--reference-doc={reference_docx}",
            f"--resource-path={ROOT_DIR / 'docs'}",
            "--standalone",
            "--toc",
            "--toc-depth=3",
        ],
    )

    _postprocess_output(output_docx, logo_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Genera la memoria DOCX con estilo URJC (portada + cabeceras).")
    parser.add_argument("--input-md", default=str(DEFAULT_INPUT_MD))
    parser.add_argument("--output-docx", default=str(DEFAULT_OUTPUT_DOCX))
    parser.add_argument("--reference-docx", default=str(DEFAULT_REFERENCE_DOCX))
    parser.add_argument("--logo", default=str(DEFAULT_LOGO))
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    build_memoria(
        input_md=Path(args.input_md),
        output_docx=Path(args.output_docx),
        reference_docx=Path(args.reference_docx),
        logo_path=Path(args.logo),
    )
    print(f"OK -> {args.output_docx}")


if __name__ == "__main__":
    main()
